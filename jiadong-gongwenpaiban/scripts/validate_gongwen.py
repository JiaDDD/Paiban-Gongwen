#!/usr/bin/env python3
"""Check that a DOCX matches the frozen 公文规格. Does not judge printed appearance."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PT_TITLE = 22.0
PT_BODY = 16.0
PT_LINE = 28.9
PT_PAGE = 14.0
MARGIN_TOP = 37.0
MARGIN_BOTTOM = 35.0
MARGIN_LEFT = 28.0
MARGIN_RIGHT = 26.0

FONT_TITLE = "方正小标宋简体"
FONT_H1 = "方正黑体简体"
FONT_H2 = "楷体_GB2312"
FONT_BODY = "方正仿宋_GB18030"
FONT_PAGE = "宋体"
FONT_LATIN = "Times New Roman"

ROLE_FONT = {
    "主标题": FONT_TITLE,
    "一级": FONT_H1,
    "二级": FONT_H2,
    "三级": FONT_BODY,
    "四级": FONT_BODY,
    "正文": FONT_BODY,
    "落款": FONT_BODY,
}

STYLE_ROLE = {
    "GW 主标题": "主标题",
    "GW 一级": "一级",
    "GW 二级": "二级",
    "GW 三级": "三级",
    "GW 四级": "四级",
    "GW 正文": "正文",
    "GW 落款": "落款",
}


def near(actual: float | None, expected: float, tolerance: float = 0.15) -> bool:
    return actual is not None and abs(actual - expected) <= tolerance


def emu_to_mm(value) -> float | None:
    if value is None:
        return None
    return value.mm


def role_of(paragraph) -> str | None:
    name = paragraph.style.name if paragraph.style is not None else ""
    return STYLE_ROLE.get(name)


def check_run_black(run, label: str, errors: list[str]) -> None:
    r_pr = run._element.rPr
    color = r_pr.find(qn("w:color")) if r_pr is not None else None
    if color is None or color.get(qn("w:val"), "").upper() != "000000":
        errors.append(f"{label} 未写成黑色")
        return
    for attr in ("themeColor", "themeTint", "themeShade"):
        if color.get(qn(f"w:{attr}")):
            errors.append(f"{label} 仍带主题色")
            return


def check_paragraph(paragraph, index: int, errors: list[str]) -> None:
    role = role_of(paragraph)
    if role is None:
        return
    label = f"第{index}段（{role}）"
    p_pr = paragraph._p.pPr
    if p_pr is None:
        errors.append(f"{label} 缺少段落属性")
        return
    widow = p_pr.find(qn("w:widowControl"))
    if widow is None or widow.get(qn("w:val"), "true") not in {"0", "false"}:
        errors.append(f"{label} 未关闭孤行控制")
    spacing = p_pr.find(qn("w:spacing"))
    line = None
    rule = None
    if spacing is not None:
        raw = spacing.get(qn("w:line"))
        line = int(raw) / 20 if raw else None
        rule = spacing.get(qn("w:lineRule"))
    if rule != "exact" or not near(line, PT_LINE, 0.05):
        errors.append(f"{label} 行距不是固定 28.9 磅")
    expected_size = PT_TITLE if role == "主标题" else PT_BODY
    east = ROLE_FONT[role]
    first_chars = 0
    ind = p_pr.find(qn("w:ind"))
    if ind is not None and ind.get(qn("w:firstLineChars")):
        first_chars = int(ind.get(qn("w:firstLineChars"))) / 100
    if role == "主标题":
        if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            errors.append(f"{label} 应对齐居中")
        if first_chars not in {0}:
            errors.append(f"{label} 不应有首行缩进")
    elif role == "落款":
        if paragraph.alignment not in {WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT}:
            errors.append(f"{label} 落款应对齐到右侧区域")
        if first_chars not in {0}:
            errors.append(f"{label} 落款不应有首行缩进")
    elif role == "正文":
        if paragraph.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            errors.append(f"{label} 正文应两端对齐")
        if not near(first_chars, 2, 0.01):
            errors.append(f"{label} 正文首行缩进应为 2 字")
    else:
        if paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT:
            errors.append(f"{label} 标题应左齐")
        if not near(first_chars, 2, 0.01):
            errors.append(f"{label} 标题首行缩进应为 2 字")
    for run in paragraph.runs:
        if not run.text:
            continue
        r_pr = run._element.rPr
        if r_pr is None or r_pr.rFonts is None:
            errors.append(f"{label} 缺少字体槽")
            break
        got_east = r_pr.rFonts.get(qn("w:eastAsia"))
        got_ascii = r_pr.rFonts.get(qn("w:ascii"))
        size = run.font.size.pt if run.font.size is not None else None
        if got_east != east:
            errors.append(f"{label} 中文字体应为 {east}，实际为 {got_east}")
            break
        if got_ascii != FONT_LATIN:
            errors.append(f"{label} 西文字体应为 {FONT_LATIN}")
            break
        if not near(size, expected_size, 0.05):
            errors.append(f"{label} 字号应为 {expected_size} 磅")
            break
        check_run_black(run, label, errors)


def check_section(document: Document, errors: list[str]) -> None:
    section = document.sections[0]
    if not near(emu_to_mm(section.page_width), 210, 0.2):
        errors.append("纸张宽度不是 A4")
    if not near(emu_to_mm(section.page_height), 297, 0.2):
        errors.append("纸张高度不是 A4")
    if not near(emu_to_mm(section.top_margin), MARGIN_TOP, 0.2):
        errors.append("上边距应为 37 mm")
    if not near(emu_to_mm(section.bottom_margin), MARGIN_BOTTOM, 0.2):
        errors.append("下边距应为 35 mm")
    if not near(emu_to_mm(section.left_margin), MARGIN_LEFT, 0.2):
        errors.append("左边距应为 28 mm")
    if not near(emu_to_mm(section.right_margin), MARGIN_RIGHT, 0.2):
        errors.append("右边距应为 26 mm")
    if not near(section.header_distance.cm if section.header_distance else None, 1.50, 0.05):
        errors.append("页眉距边界应为 1.50 cm")
    if not near(section.footer_distance.cm if section.footer_distance else None, 2.50, 0.05):
        errors.append("页脚距边界应为 2.50 cm")
    sect_pr = section._sectPr
    grid = sect_pr.find(qn("w:docGrid"))
    if grid is None:
        errors.append("缺少指定行和字符网格")
    else:
        if grid.get(qn("w:type")) != "linesAndChars":
            errors.append("文档网格类型应为 linesAndChars")
        pitch = grid.get(qn("w:linePitch"))
        if pitch is None or abs(int(pitch) / 20 - PT_LINE) > 0.1:
            errors.append("网格行距应为 28.9 磅")
    if sect_pr.find(qn("w:evenAndOddHeaders")) is None:
        errors.append("未开启奇偶页不同页眉页脚")


def check_page_numbers(document: Document, errors: list[str]) -> None:
    section = document.sections[0]
    for name, footer in (("奇数页脚", section.footer), ("偶数页脚", section.even_page_footer)):
        texts = []
        has_page = False
        for paragraph in footer.paragraphs:
            texts.append(paragraph.text)
            xml = paragraph._p.xml
            if "PAGE" in xml:
                has_page = True
            for run in paragraph.runs:
                r_pr = run._element.rPr
                if r_pr is None or r_pr.rFonts is None:
                    continue
                if r_pr.rFonts.get(qn("w:eastAsia")) != FONT_PAGE:
                    errors.append(f"{name} 页码中文字体应为宋体")
                if r_pr.rFonts.get(qn("w:ascii")) != FONT_PAGE:
                    errors.append(f"{name} 页码西文槽应为宋体")
        if not has_page:
            errors.append(f"{name} 缺少 PAGE 域")
        joined = "".join(texts)
        if "\u2014" not in joined:
            errors.append(f"{name} 页码装饰应为 \u2014 n \u2014")


def validate(path: Path) -> list[str]:
    document = Document(str(path))
    errors: list[str] = []
    check_section(document, errors)
    check_page_numbers(document, errors)
    managed = 0
    for index, paragraph in enumerate(document.paragraphs, start=1):
        if role_of(paragraph):
            managed += 1
            check_paragraph(paragraph, index, errors)
    if managed == 0:
        errors.append("未找到 GW 命名样式段落，无法确认规格已套用")
    seen: set[str] = set()
    unique: list[str] = []
    for item in errors:
        if item not in seen:
            seen.add(item)
            unique.append(item)
    return unique


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    args = parser.parse_args()
    path = Path(args.input)
    if not path.exists():
        print(f"VALIDATE\tFAIL\tmissing {path}", file=sys.stderr)
        return 1
    errors = validate(path)
    if errors:
        print("VALIDATE\tFAIL")
        for item in errors:
            print(f"ERROR\t{item}")
        return 1
    print("VALIDATE\tOK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
