#!/usr/bin/env python3
"""Create or reformat a .docx using the frozen 公文排版 spec."""

from __future__ import annotations

import argparse
import math
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

PT_TITLE = 22.0
PT_BODY = 16.0
PT_LINE = 28.9
PT_PAGE = 14.0

FONT_TITLE = "方正小标宋简体"
FONT_H1 = "方正黑体简体"
FONT_H2 = "楷体_GB2312"
FONT_BODY = "方正仿宋_GB18030"
FONT_PAGE = "宋体"
FONT_LATIN = "Times New Roman"

CN_NUM = "零一二三四五六七八九十"

H1_PREFIX = re.compile(r"^[一二三四五六七八九十百千]+、")
H2_PREFIX = re.compile(r"^（[一二三四五六七八九十百千]+）")
H3_PREFIX = re.compile(r"^\d+\.")
H4_PREFIX = re.compile(r"^（\d+）")
ATX = re.compile(r"^(#{1,5})\s+(.*)$")
IMAGE = re.compile(r"^!\[.*?\]\(.*\)$")
LINK = re.compile(r"\[([^\]]+)\]\([^)]+\)")
MD_WRAP = re.compile(r"(\*\*|__|\*|_)(.*?)\1")
LATIN_CHUNK = re.compile(r"[A-Za-z0-9][A-Za-z0-9+\-._/%]*")
LIST_ITEM = re.compile(r"^(?:\d+\.|[-*+]|（[一二三四五六七八九十百\d]+）|\([0-9]+\))\s+")
DATE_RE = re.compile(r"^\d{4}年\d{1,2}月\d{1,2}日(?:\s*印发)?[。．.]?$")
OFFICE_RE = re.compile(r"(?:单位|办公室|委员会|人民政府|政府|局|厅|部|院|中心|公司|集团|学校|协会|研究院)$")

ROLE_STYLE = {
    "主标题": "GW 主标题",
    "一级": "GW 一级",
    "二级": "GW 二级",
    "三级": "GW 三级",
    "四级": "GW 四级",
    "正文": "GW 正文",
    "落款": "GW 落款",
}
ROLE_FONT = {
    "主标题": FONT_TITLE,
    "一级": FONT_H1,
    "二级": FONT_H2,
    "三级": FONT_BODY,
    "四级": FONT_BODY,
    "正文": FONT_BODY,
    "落款": FONT_BODY,
}


def cn_num(n: int) -> str:
    if n <= 10:
        return CN_NUM[n]
    if n < 20:
        return "十" + (CN_NUM[n - 10] if n > 10 else "")
    if n < 100:
        tens, ones = divmod(n, 10)
        return CN_NUM[tens] + "十" + (CN_NUM[ones] if ones else "")
    return str(n)


def strip_md(text: str) -> str:
    text = text.strip()
    text = LINK.sub(r"\1", text)
    text = MD_WRAP.sub(r"\2", text)
    text = text.replace("`", "")
    if text.startswith(">"):
        text = text.lstrip(">").strip()
    if re.match(r"^[-*+]\s+", text):
        text = re.sub(r"^[-*+]\s+", "", text)
    return re.sub(r"[ \t]+", " ", text).strip()


def parse_markdown(text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    h1 = h2 = h3 = h4 = 0
    pending: list[str] = []

    def flush() -> None:
        nonlocal pending
        joined = strip_md("".join(pending).strip())
        pending = []
        if joined:
            blocks.append(("正文", joined))

    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped:
            flush()
            i += 1
            continue
        if IMAGE.match(stripped):
            flush()
            i += 1
            continue
        if stripped in {"[落款]", "【落款】"}:
            flush()
            unit = date = ""
            if i + 1 < len(lines) and lines[i + 1].strip():
                unit = strip_md(lines[i + 1])
                i += 1
            if i + 1 < len(lines) and lines[i + 1].strip():
                date = strip_md(lines[i + 1])
                i += 1
            if unit:
                blocks.append(("落款", unit))
            if date:
                blocks.append(("落款", date))
            i += 1
            continue
        if LIST_ITEM.match(stripped):
            flush()
            item = strip_md(stripped)
            if item:
                blocks.append(("正文", item))
            i += 1
            continue
        m = ATX.match(stripped)
        if m:
            flush()
            level = len(m.group(1))
            title = strip_md(m.group(2))
            if level == 1:
                blocks.append(("主标题", title))
            elif level == 2:
                if not H1_PREFIX.match(title):
                    h1 += 1
                    h2 = h3 = h4 = 0
                    title = f"{cn_num(h1)}、{title}"
                else:
                    h1 += 1
                    h2 = h3 = h4 = 0
                blocks.append(("一级", title))
            elif level == 3:
                if not H2_PREFIX.match(title):
                    h2 += 1
                    h3 = h4 = 0
                    title = f"（{cn_num(h2)}）{title}"
                else:
                    h2 += 1
                    h3 = h4 = 0
                blocks.append(("二级", title))
            elif level == 4:
                if not H3_PREFIX.match(title):
                    h3 += 1
                    h4 = 0
                    title = f"{h3}.{title}"
                else:
                    h3 += 1
                    h4 = 0
                blocks.append(("三级", title))
            else:
                if not H4_PREFIX.match(title):
                    h4 += 1
                    title = f"（{h4}）{title}"
                else:
                    h4 += 1
                blocks.append(("四级", title))
            i += 1
            continue
        pending.append(stripped)
        i += 1
    flush()

    if blocks and blocks[-1][0] == "正文" and re.fullmatch(r"[—–-]{1,2}.{1,20}", blocks[-1][1]):
        blocks[-1] = ("落款", blocks[-1][1])
    return blocks


def set_run_black(run) -> None:
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    color = r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    color.set(qn("w:val"), "000000")
    for attr in ("themeColor", "themeTint", "themeShade"):
        color.attrib.pop(qn(f"w:{attr}"), None)


def set_run_fonts(run, east_asia: str, latin: str, size_pt: float, bold: bool = False) -> None:
    run.bold = bold
    run.italic = False
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:cs"), latin)
    for tag in ("w:sz", "w:szCs"):
        node = r_pr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            r_pr.append(node)
        node.set(qn("w:val"), str(int(round(size_pt * 2))))
    set_run_black(run)


def add_mixed_runs(paragraph, text: str, east_asia: str, size_pt: float) -> None:
    idx = 0
    for match in LATIN_CHUNK.finditer(text):
        if match.start() > idx:
            run = paragraph.add_run(text[idx : match.start()])
            set_run_fonts(run, east_asia, FONT_LATIN, size_pt)
        run = paragraph.add_run(match.group(0))
        set_run_fonts(run, east_asia, FONT_LATIN, size_pt)
        idx = match.end()
    if idx < len(text):
        run = paragraph.add_run(text[idx:])
        set_run_fonts(run, east_asia, FONT_LATIN, size_pt)
    if not text:
        run = paragraph.add_run("")
        set_run_fonts(run, east_asia, FONT_LATIN, size_pt)


def set_widow_off(paragraph) -> None:
    paragraph.paragraph_format.widow_control = False
    p_pr = paragraph._p.get_or_add_pPr()
    widow = p_pr.find(qn("w:widowControl"))
    if widow is None:
        widow = OxmlElement("w:widowControl")
        p_pr.append(widow)
    widow.set(qn("w:val"), "0")


def set_paragraph_geometry(paragraph, role: str, signature_left: int = 0) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(PT_LINE)
    set_widow_off(paragraph)

    p_pr = paragraph._p.get_or_add_pPr()
    snap = p_pr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        p_pr.append(snap)

    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)

    indent = p_pr.find(qn("w:ind"))
    if indent is None:
        indent = OxmlElement("w:ind")
        p_pr.append(indent)

    if role == "主标题":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        jc.set(qn("w:val"), "center")
        for attr in list(indent.attrib):
            indent.attrib.pop(attr)
    elif role == "落款":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        jc.set(qn("w:val"), "left")
        for attr in list(indent.attrib):
            indent.attrib.pop(attr)
        indent.set(qn("w:left"), str(signature_left))
        indent.set(qn("w:firstLine"), "0")
        indent.set(qn("w:firstLineChars"), "0")
    elif role == "正文":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        jc.set(qn("w:val"), "both")
        indent.set(qn("w:firstLineChars"), "200")
        indent.set(qn("w:firstLine"), "0")
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        jc.set(qn("w:val"), "left")
        indent.set(qn("w:firstLineChars"), "200")
        indent.set(qn("w:firstLine"), "0")


def signature_left_twips(lines: list[str]) -> int:
    usable = Mm(210 - 28 - 26).twips

    def width_em(text: str) -> float:
        return sum(
            0 if unicodedata.category(char) in {"Mn", "Me", "Cf"} else 4 if char == "\t" else 1
            for char in text.strip()
        )

    longest = max((width_em(line) for line in lines if line.strip()), default=0)
    headroom = max(2.0, longest * 0.15)
    width = math.ceil((longest + headroom) * PT_BODY * 20)
    return max(0, int(usable - width))


def ensure_styles(document: Document) -> None:
    for role, style_name in ROLE_STYLE.items():
        try:
            style = document.styles[style_name]
        except KeyError:
            style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        east = ROLE_FONT[role]
        size = PT_TITLE if role == "主标题" else PT_BODY
        style.font.name = FONT_LATIN
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn("w:eastAsia"), east)
        r_fonts.set(qn("w:ascii"), FONT_LATIN)
        r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
        r_fonts.set(qn("w:cs"), FONT_LATIN)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(PT_LINE)
        style.paragraph_format.widow_control = False


def add_paragraph(document: Document, role: str, text: str, signature_left: int = 0):
    p = document.add_paragraph()
    p.style = ROLE_STYLE[role]
    set_paragraph_geometry(p, role, signature_left)
    east = ROLE_FONT[role]
    size = PT_TITLE if role == "主标题" else PT_BODY
    add_mixed_runs(p, text, east, size)
    return p


def configure_section(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)
    section.gutter = Mm(0)
    section.header_distance = Cm(1.50)
    section.footer_distance = Cm(2.50)
    section.different_first_page_header_footer = False
    sect_pr = section._sectPr
    even = sect_pr.find(qn("w:evenAndOddHeaders"))
    if even is None:
        even = OxmlElement("w:evenAndOddHeaders")
        sect_pr.append(even)
    grid = sect_pr.find(qn("w:docGrid"))
    if grid is None:
        grid = OxmlElement("w:docGrid")
        sect_pr.append(grid)
    grid.set(qn("w:type"), "linesAndChars")
    grid.set(qn("w:linePitch"), str(int(round(PT_LINE * 20))))
    grid.set(qn("w:charSpace"), "0")


def _page_run_rpr() -> OxmlElement:
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), FONT_PAGE)
    r_pr.append(r_fonts)
    hint = OxmlElement("w:hint")
    hint.set(qn("w:val"), "eastAsia")
    r_pr.append(hint)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)
    bold = OxmlElement("w:b")
    bold.set(qn("w:val"), "0")
    r_pr.append(bold)
    for tag in ("w:sz", "w:szCs"):
        node = OxmlElement(tag)
        node.set(qn("w:val"), str(int(PT_PAGE * 2)))
        r_pr.append(node)
    return r_pr


def add_page_number_paragraph(paragraph, align: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = None
    set_widow_off(paragraph)
    p_pr_align = paragraph._p.get_or_add_pPr()
    jc = p_pr_align.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr_align.append(jc)
    jc.set(qn("w:val"), "center")

    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:rPr"))
    if existing is not None:
        p_pr.remove(existing)
    p_pr.append(_page_run_rpr())

    def add_text(text: str) -> None:
        run = paragraph.add_run(text)
        run._element.append(_page_run_rpr())
        set_run_fonts(run, FONT_PAGE, FONT_PAGE, PT_PAGE)

    add_text("— ")
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)
    begin_run._r.append(_page_run_rpr())
    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE \\* CHARFORMAT "
    instr_run._r.append(instr)
    instr_run._r.append(_page_run_rpr())
    sep_run = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(sep)
    sep_run._r.append(_page_run_rpr())
    result = paragraph.add_run("1")
    set_run_fonts(result, FONT_PAGE, FONT_PAGE, PT_PAGE)
    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    end_run._r.append(_page_run_rpr())
    add_text(" —")


def configure_footers(section) -> None:
    for footer, align in ((section.footer, "center"), (section.even_page_footer, "center")):
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            p = footer.paragraphs[0]
            p.clear()
        else:
            p = footer.add_paragraph()
        add_page_number_paragraph(p, align)
        for extra in footer.paragraphs[1:]:
            extra._element.getparent().remove(extra._element)
    for header in (section.header, section.even_page_header):
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.text = ""


def apply_document_defaults(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(PT_BODY)
    r_pr = normal.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), FONT_BODY)
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:cs"), FONT_LATIN)


def build_document(blocks: list[tuple[str, str]]) -> Document:
    document = Document()
    apply_document_defaults(document)
    ensure_styles(document)
    section = document.sections[0]
    configure_section(section)
    configure_footers(section)

    body = document.element.body
    for child in list(body):
        if child.tag == qn("w:p"):
            body.remove(child)

    sig_lines = [text for role, text in blocks if role == "落款" and text.strip()]
    sig_left = signature_left_twips(sig_lines)

    title_seen = False
    signature_gap = False
    for role, text in blocks:
        if role == "主标题":
            add_paragraph(document, "主标题", text)
            add_paragraph(document, "正文", "")
            title_seen = True
            continue
        if role == "落款" and not signature_gap:
            add_paragraph(document, "正文", "")
            signature_gap = True
        add_paragraph(document, role, text, sig_left if role == "落款" else 0)
    if not title_seen and blocks:
        pass
    return document


def unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    stem, suffix = path.stem, path.suffix
    n = 2
    while True:
        candidate = path.with_name(f"{stem}-{n}{suffix}")
        if not candidate.exists():
            return candidate
        n += 1


def classify_docx(path: Path) -> list[tuple[str, str]]:
    document = Document(str(path))
    blocks: list[tuple[str, str]] = []
    for p in document.paragraphs:
        text = strip_md(p.text)
        if not text:
            continue
        style = (p.style.name or "") if p.style is not None else ""
        style_l = style.lower()
        if style in ROLE_STYLE.values():
            role = next(role for role, name in ROLE_STYLE.items() if name == style)
            blocks.append((role, text))
        elif "title" in style_l or style == "Title":
            blocks.append(("主标题", text))
        elif style_l in {"heading 1", "heading1"} or H1_PREFIX.match(text):
            blocks.append(("一级", text))
        elif style_l in {"heading 2", "heading2"} or H2_PREFIX.match(text):
            blocks.append(("二级", text))
        elif style_l in {"heading 4", "heading4"} or H4_PREFIX.match(text):
            blocks.append(("四级", text))
        elif style_l in {"heading 3", "heading3"}:
            blocks.append(("三级", text))
        elif DATE_RE.match(text) or re.fullmatch(r"[—–-]{1,2}.{1,20}", text) or OFFICE_RE.search(text):
            blocks.append(("落款", text))
        else:
            blocks.append(("正文", text))
    rebuilt: list[tuple[str, str]] = []
    h1 = h2 = h3 = h4 = 0
    for role, text in blocks:
        if role == "一级" and not H1_PREFIX.match(text):
            h1 += 1
            h2 = h3 = h4 = 0
            text = f"{cn_num(h1)}、{text}"
        elif role == "二级" and not H2_PREFIX.match(text):
            h2 += 1
            h3 = h4 = 0
            text = f"（{cn_num(h2)}）{text}"
        elif role == "三级" and not H3_PREFIX.match(text):
            h3 += 1
            h4 = 0
            text = f"{h3}.{text}"
        elif role == "四级" and not H4_PREFIX.match(text):
            h4 += 1
            text = f"（{h4}）{text}"
        rebuilt.append((role, text))
    return rebuilt


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("command", choices=["create", "format"])
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    src = Path(args.input)
    dest = unique_path(Path(args.output))
    dest.parent.mkdir(parents=True, exist_ok=True)

    if args.command == "create":
        text = src.read_text(encoding="utf-8")
        blocks = parse_markdown(text)
    else:
        blocks = classify_docx(src)

    if not blocks:
        print("No usable blocks", file=sys.stderr)
        return 1

    for role, text in blocks:
        shown = text if text else "⟨空行⟩"
        print(f"{role}\t{shown}")

    document = build_document(blocks)
    document.save(str(dest))
    print(f"DOCX\t{dest}")

    validator = Path(__file__).resolve().parent / "validate_gongwen.py"
    if validator.exists():
        from validate_gongwen import validate

        errors = validate(dest)
        if errors:
            print("VALIDATE\tFAIL")
            for item in errors:
                print(f"ERROR\t{item}")
            return 1
        print("VALIDATE\tOK")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
